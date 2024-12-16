from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def create_document():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    return doc

def set_column_widths_oxml(table, widths):
    tbl = table._tbl  # Access the underlying <w:tbl> element
    tbl_grid = OxmlElement('w:tblGrid')

    for width in widths:
        grid_col = OxmlElement('w:gridCol')
        grid_col.set(qn('w:w'), str(width))
        tbl_grid.append(grid_col)
    
    tbl.insert(0, tbl_grid)

def add_subscription_table(doc, data):
    subscriptions = data['Objects']
    
    # Add 'Subscriptions' as a heading
    doc.add_heading('Subscriptions', level=1)

    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Subscription ID'
    hdr_cells[1].text = 'Name'
    for subscription in subscriptions:
        row_cells = table.add_row().cells
        row_cells[0].text = subscription
        if 'name' in subscriptions[subscription]:
            row_cells[1].text = subscriptions[subscription]['name']
        else:
            row_cells[1].text = 'Name not found'

        # Reduce paragraph spacing
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)

def add_resource_group_table(doc, data):
    objects = data['Objects']

    # Add 'Resource Groups' as a heading
    doc.add_heading('Resource Groups', level=1)

    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Subscription ID'
    hdr_cells[1].text = 'Resource Group'

    resource_groups = set()

    for subscription_id in objects:
        for user_name, resources in objects[subscription_id].items():
            for resource in resources:
                details = resource['Details']
                if isinstance(details, dict):
                    # If Details is a dict, proceed as before
                    if 'id' in details:
                        match = re.search(r'/resourceGroups/(.*?)/', details['id'])
                        if match:
                            resource_group = match.group(1)
                            resource_groups.add((subscription_id, resource_group))
                elif isinstance(details, list):
                    # If Details is a list, iterate over each item
                    for item in details:
                        if 'id' in item:
                            match = re.search(r'/resourceGroups/(.*?)/', item['id'])
                            if match:
                                resource_group = match.group(1)
                                resource_groups.add((subscription_id, resource_group))
                else:
                    # Handle any other cases if necessary
                    pass

    for subscription_id, rg in resource_groups:
        row_cells = table.add_row().cells
        row_cells[0].text = subscription_id
        row_cells[1].text = rg

        # Reduce paragraph spacing
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)

    doc.add_paragraph('\n')


def group_resources_by_resource_group(resources):
    resource_groups = {}
    for resource in resources:
        details = resource['Details']
        if isinstance(details, dict):
            # If Details is a dict
            if 'id' in details:
                match = re.search(r'/resourceGroups/(.*?)/', details['id'])
                if match:
                    resource_group = match.group(1)
                    if resource_group not in resource_groups:
                        resource_groups[resource_group] = []
                    resource_groups[resource_group].append(resource)
        elif isinstance(details, list):
            # If Details is a list
            for item in details:
                if 'id' in item:
                    match = re.search(r'/resourceGroups/(.*?)/', item['id'])
                    if match:
                        resource_group = match.group(1)
                        if resource_group not in resource_groups:
                            resource_groups[resource_group] = []
                        resource_groups[resource_group].append(resource)
        else:
            # Handle any other cases if necessary
            pass
    return resource_groups
def sort_resources_by_type(resources):
    return sorted(resources, key=lambda x: x['ResourceType'])

def format_attribute(attribute):
    attribute = attribute.replace('_', ' ')
    return attribute.capitalize()

def add_resource_table(doc, resource):
    if not resource['Details']:
        print(f"No details found for resource: {resource}")
        return
    
    details = resource['Details']
    recommendations = resource.get('Recommendations', [])
    
    # Determine the resource name
    if isinstance(details, dict):
        resource_name = details.get('name', 'Unknown')
    elif isinstance(details, list):
        # If Details is a list, try to get the name from the first item
        if details and 'name' in details[0]:
            resource_name = details[0]['name']
        else:
            resource_name = 'Unknown'
    else:
        resource_name = 'Unknown'
    
    resource_type = format_resource_type(resource['ResourceType'])
    add_heading(doc, resource_type, resource_name)

    # Add resource details
    add_dict_as_table(doc, details, 3)

    # Add recommendations if available
    if recommendations:
        doc.add_heading('Recommendations', level=4)
        for recommendation in recommendations:
            add_dict_as_table(doc, recommendation, 4)

    doc.add_paragraph()

def format_resource_type(resource_type):
    # Remove 'Microsoft.' from the resource type
    return resource_type.replace('Microsoft.', '')

def add_heading(doc, resource_type, resource_name, level=3):
    # Add resource type and name as a heading
    doc.add_heading(f"{resource_type} - {resource_name}", level)

def add_dict_as_table(doc, value, level):
    if isinstance(value, dict):
        table = doc.add_table(rows=0, cols=2)
        set_column_widths_oxml(table, [3500, 250])
        for attr, val in value.items():
            if isinstance(val, (dict, list)):
                doc.add_heading(format_attribute(attr), level+1)
                add_dict_as_table(doc, val, level+1)
            else:
                add_row_to_table(table, attr, val)
    elif isinstance(value, list):
        for index, item in enumerate(value):
            doc.add_heading(f"Item {index+1}", level+1)
            add_dict_as_table(doc, item, level+1)
    else:
        # Handle other data types if necessary
        pass

def add_row_to_table(table, attr, val):
    row_cells = table.add_row().cells
    row_cells[0].text = format_attribute(attr)
    row_cells[1].text = str(val)
    reduce_paragraph_spacing(row_cells)

def reduce_paragraph_spacing(row_cells):
    # Reduce paragraph spacing
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)

def add_resource_tables(doc, data):
    objects = data['Objects']
    for subscription_id in objects:
        for user_name, resources in objects[subscription_id].items():
            resource_groups = group_resources_by_resource_group(resources)
            for resource_group, resources in resource_groups.items():
                doc.add_heading("Resource Group: " + resource_group, level=2)
                sorted_resources = sort_resources_by_type(resources)
                for resource in sorted_resources:
                    add_resource_table(doc, resource)

def generate_word_document(data, output_folder):
    doc = create_document()
    add_subscription_table(doc, data)
    add_resource_group_table(doc, data)
    add_resource_tables(doc, data)
    doc.save(f'{output_folder}/output.docx')